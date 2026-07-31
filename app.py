import streamlit as st
from docxtpl import DocxTemplate
import docx
import io
import os
from datetime import date

# Configuração visual profissional para smartphone
st.set_page_config(page_title="Gerador ETP Oficial", page_icon="📝", layout="centered")

st.markdown("<h2 style='text-align: center; color: #0284c7;'>📝 Preenchedor de ETP Digital BOA ESPERANÇA - MG</h2>", unsafe_allow_html=True)
st.write("Insira as informações abaixo para preencher o modelo oficial de Boa Esperança/MG:")

# --- FORMULÁRIO DIVIDIDO POR SEÇÕES ---
st.subheader("📌 Informações Básicas")
secretaria = st.text_input("Secretaria demandante", value="Diretoria de Trânsito e Sinalização Pública")
objeto = st.text_area("Objeto da Contratação", placeholder="Ex: Aquisição de materiais de sinalização...")

st.subheader("📄 Itens do Formulário")
desc_necessidade = st.text_area("1 - Descrição da necessidade da contratação")
previsao_pac = st.text_area("2 - Previsão da contratação no Plano Anual de Contratações")
requisitos = st.text_area("3.1 - Requisitos da contratação")
inicio_servicos = st.text_area("3.2 - Prazo de implantação e início dos serviços")
tipo_prestacao = st.text_area("3.3 - Da prestação dos serviços")
fiscalizacao = st.text_area("3.4 - Da fiscalização e acompanhamento")
garantia = st.text_area("3.5 - Da garantia e níveis de serviço")
pagamento = st.text_area("3.6 - Do pagamento")
estimativas_quant = st.text_area("4 - Estimativas das quantidades para contratação")

st.subheader("🛒 Tabela de Itens (Demonstração)")
st.info("O robô irá gerar uma linha com os dados abaixo para validar o modelo.")
item_num = st.text_input("Número do Item", "01")
item_desc = st.text_input("Descrição do Item", "Material de Exemplo")
item_unid = st.text_input("Unidade de Medida", "UNID")
item_quant = st.number_input("Quantidade", min_value=1, value=1)
item_total = st.number_input("Total", min_value=1, value=1)

st.subheader("📊 Mercado e Valores")
levantamento = st.text_area("5 - Levantamento de mercado")
estimativa_valor = st.text_area("6 - Estimativa do valor da contratação")
desc_solucao = st.text_area("7 - Descrição da solução como um todo")
justificativa_parc = st.text_area("8 - Justificativa do parcelamento ou não")
demonstrativo_res = st.text_area("9 - Demonstrativo dos resultados pretendidos")
providencias_previas = st.text_area("10 - Providências a serem adotadas previamente")
contratacoes_corr = st.text_area("11 - Contratações correlatas e/ou interdependentes")
impactos_amb = st.text_area("12 - Descrição de possíveis impactos ambientais")
conclusao = st.text_area("13 - Conclusão")

# --- CONSTRUÇÃO DO DICIONÁRIO DE DADOS ---
dados_etp = {
    "secretaria_demandante2": secretaria,
    "objeto": objeto,
    "descricao_necessidade": desc_necessidade,
    "previsao_pac": previsao_pac,
    "requisitos_contratacao": requisitos,
    "inicio_servicos": inicio_servicos,
    "tipo_prestacao": tipo_prestacao,
    "fiscalizacao_acompanhamento": fiscalizacao,
    "garantia_niveis_servico": garantia,
    "pagamento": pagamento,
    "estimativas_contrataçao": estimativas_quant,
    "itens": [
        {
            "item": item_num,
            "descricao": item_desc,
            "unidade": item_unid,
            "quantidade": item_quant,
            "total": item_total
        }
    ],
    "levantamento_mercado": levantamento,
    "estimativa_valor": estimativa_valor,
    "descricao_solucao": desc_solucao,
    "justificativa_parcelamento": justificativa_parc,
    "demonstrativo_resultados": demonstrativo_res,
    "providencias_previas": providencias_previas,
    "contratacoes_correlatas": contratacoes_corr,
    "impactos_ambientais": impactos_amb,
    "conclusao": conclusao,
    "data": date.today().strftime("%d/%m/%Y")
}

st.markdown("---")

def substituir_no_cabecalho(paragraph, texto_substituto):
    """Une fragmentos quebrados pelo Word e força a substituição da tag"""
    texto_completo = "".join([run.text for run in paragraph.runs])
    
    # Mapeia variações de escrita com e sem espaços internos
    alvos = ['{{ secretaria_demandante1 }}', '{{secretaria_demandante1}}']
    
    substituiu = False
    for alvo in alvos:
        if alvo in texto_completo:
            texto_completo = texto_completo.replace(alvo, texto_substituto)
            substituiu = True
            
    if substituiu:
        # Mantém a formatação do primeiro fragmento se houver, limpa o resto e injeta o texto correto
        if len(paragraph.runs) > 0:
            p_run = paragraph.runs[0]
            p_run.text = texto_completo
            # Apaga os fragmentos fantasmas que quebravam a lógica
            for r in paragraph.runs[1:]:
                r.text = ""
        else:
            paragraph.text = texto_completo

# --- PROCESSAMENTO DO MODELO ---
if st.button("🚀 GERAR DOCUMENTO ETP OFICIAL"):
    if not objeto or not desc_necessidade:
        st.error("❌ Preencha pelo menos o Objeto e a Descrição da Necessidade para testar.")
    else:
        with st.spinner("Processando ETP na nuvem..."):
            try:
                if not os.path.exists("modelo_etp.docx"):
                    st.error("❌ Erro: O arquivo 'modelo_etp.docx' não foi encontrado no servidor.")
                else:
                    # 1. Abre e preenche o corpo normal com DocxTemplate (Substitui secretaria_demandante2 e objeto)
                    doc_tpl = DocxTemplate("modelo_etp.docx")
                    doc_tpl.render(dados_etp)
                    
                    # Salva temporariamente em memória para manipular o cabeçalho
                    buffer_intermediario = io.BytesIO()
                    doc_tpl.save(buffer_intermediario)
                    buffer_intermediario.seek(0)
                    
                    # 2. Abre usando python-docx tradicional para forçar a varredura do Cabeçalho
                    doc_final = docx.Document(buffer_intermediario)
                    
                    # Percorre todas as seções e tabelas de cabeçalho aplicando o reconstrutor de parágrafos
                    for section in doc_final.sections:
                        header = section.header
                        if header is not None:
                            # Corrige parágrafos soltos no cabeçalho
                            for paragraph in header.paragraphs:
                                substituir_no_cabecalho(paragraph, secretaria)
                            
                            # Corrige tabelas dentro do cabeçalho (caso o timbre use tabelas estruturais)
                            for table in header.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        for paragraph in cell.paragraphs:
                                            substituir_no_cabecalho(paragraph, secretaria)
                    
                    # Envia o arquivo preenchido final para a memória de download
                    buffer = io.BytesIO()
                    doc_final.save(buffer)
                    buffer.seek(0)
                    
                    st.success("✅ ETP estruturado com sucesso!")
                    
                    st.download_button(
                        label="📥 BAIXAR ETP PREENCHIDO (.DOCX)",
                        data=buffer,
                        file_name=f"ETP_{secretaria[:15]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            except Exception as e:
                st.error(f"❌ Ocorreu um erro ao processar o documento: {e}")
                                                     
